from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd
import datetime

class ReportGenerator:
    def __init__(self):
        self.doc = Document()
        
    def generate_report(self, df, selected_item, stats_df, screening_df, figures, model_metrics):
        """
        Generates a Word report.
        
        Args:
            df (pd.DataFrame): The full dataset.
            selected_item (str): The item being analyzed.
            stats_df (pd.DataFrame): Basic statistics for the item.
            screening_df (pd.DataFrame): Full screening results.
            figures (dict): Dictionary of {name: BytesIO} for images.
            model_metrics (dict): RMSE metrics for the selected item.
        """
        # --- Title Page ---
        self.doc.add_heading('미사일 신뢰도 분석 결과 보고서', 0)
        
        p = self.doc.add_paragraph()
        p.add_run(f'생성 일시: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n').bold = True
        p.add_run(f'분석 대상 항목: Item {selected_item}').bold = True
        
        self.doc.add_page_break()
        
        # --- 1. Data Summary ---
        self.doc.add_heading('1. 데이터 개요 (Data Summary)', level=1)
        
        # Count Summary
        counts = df['Dataset'].value_counts()
        self.doc.add_paragraph('데이터셋 구성:')
        for dataset, count in counts.items():
            self.doc.add_paragraph(f'- {dataset}: {count} 개', style='List Bullet')
            
        # Time Range
        self.doc.add_paragraph('\n운용월 범위:')
        for dataset in df['Dataset'].unique():
            subset = df[df['Dataset'] == dataset]
            min_month = subset['운용월'].min()
            max_month = subset['운용월'].max()
            self.doc.add_paragraph(f'- {dataset}: {min_month} ~ {max_month} 개월', style='List Bullet')

        # --- 2. Detailed Analysis ---
        self.doc.add_heading(f'2. 상세 분석: Item {selected_item}', level=1)
        
        # 2.1 Statistics
        self.doc.add_heading('2.1 기초 통계량', level=2)
        self._add_dataframe_to_table(stats_df)
        
        # 2.2 Distribution
        self.doc.add_heading('2.2 데이터 분포 (Distribution)', level=2)
        if 'distribution_plot' in figures:
            self.doc.add_picture(figures['distribution_plot'], width=Inches(6))
            self.doc.add_paragraph('그림 1. 데이터 분포 비교 (KDE)', style='Caption')
            
        if 'box_plot' in figures:
            self.doc.add_picture(figures['box_plot'], width=Inches(6))
            self.doc.add_paragraph('그림 2. 데이터 분포 요약 (Box Plot)', style='Caption')
            
        # 2.3 Trend Prediction
        self.doc.add_heading('2.3 추세 예측 (Trend Prediction)', level=2)
        if 'trend_plot' in figures:
            self.doc.add_picture(figures['trend_plot'], width=Inches(6))
            self.doc.add_paragraph('그림 3. 다중 모델 추세 예측 결과', style='Caption')
            
        # RMSE Table
        self.doc.add_heading('모델 성능 평가 (RMSE)', level=3)
        if model_metrics:
            metrics_df = pd.DataFrame(list(model_metrics.items()), columns=['Model', 'RMSE'])
            metrics_df = metrics_df.sort_values('RMSE')
            self._add_dataframe_to_table(metrics_df)
            
            best_model = metrics_df.iloc[0]
            self.doc.add_paragraph(f'\n해석: 현재 데이터에 대해 가장 적합한 모델은 "{best_model["Model"]}"이며, RMSE는 {best_model["RMSE"]:.4f}입니다.')
        
        self.doc.add_page_break()
        
        # --- 3. Screening Summary ---
        self.doc.add_heading('3. 전체 항목 스크리닝 (Screening Summary)', level=1)
        self.doc.add_paragraph('전체 27개 항목에 대한 위험도 분석 결과 (Top 10):')
        
        top_10 = screening_df.head(10)
        # Select relevant columns for report
        cols_to_show = ['Item', 'Slope', 'R2', 'Var_Ratio', 'Norm_Slope']
        # Check if columns exist (v1 code might have different names)
        available_cols = [c for c in cols_to_show if c in top_10.columns]
        self._add_dataframe_to_table(top_10[available_cols])
        
        return self._save_to_stream()

    def _add_dataframe_to_table(self, df):
        """Helper to convert DataFrame to Word Table."""
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            
        # Rows
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)
                    
    def _save_to_stream(self):
        """Saves the document to a BytesIO stream."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
