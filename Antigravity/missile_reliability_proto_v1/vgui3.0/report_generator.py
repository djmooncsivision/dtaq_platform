from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd
import datetime

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class ReportGenerator:
    def __init__(self):
        self.doc = Document()
        
    def generate_report(self, df, items_data, screening_df, is_full_report=False):
        """
        Generates a Word report.
        
        Args:
            df (pd.DataFrame): The full dataset.
            items_data (list or dict): 
                If Single Report: {'item': str, 'stats': df, 'figures': dict, 'metrics': dict}
                If Full Report: List of dicts [{'item': str, ...}, ...]
            screening_df (pd.DataFrame): Full screening results.
            is_full_report (bool): Whether to generate for all items.
        """
        # --- Title Page ---
        title = '미사일 신뢰도 분석 결과 보고서 (전체)' if is_full_report else f'미사일 신뢰도 분석 결과 보고서 (Item {items_data["item"]})'
        self.doc.add_heading(title, 0)
        
        p = self.doc.add_paragraph()
        p.add_run(f'생성 일시: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n').bold = True
        if not is_full_report:
            p.add_run(f'분석 대상 항목: Item {items_data["item"]}').bold = True
        else:
            p.add_run(f'분석 대상: 전체 {len(items_data)}개 항목').bold = True
        
        self.doc.add_page_break()
        
        # --- Table of Contents ---
        self.doc.add_heading('목차 (Table of Contents)', level=1)
        self._add_toc()
        self.doc.add_page_break()
        
        # --- 1. Data Summary ---
        self.doc.add_heading('1. 데이터 개요 (Data Summary)', level=1)
        
        # Count Summary
        counts = df['Dataset'].value_counts()
        self.doc.add_paragraph('데이터셋 구성:')
        for dataset, count in counts.items():
            self.doc.add_paragraph(f'- {dataset}: {count} 개', style='List Bullet')
            
        # Time Range
        self.doc.add_paragraph('\n운용월 범위:')
        for dataset in df['Dataset'].unique():
            subset = df[df['Dataset'] == dataset]
            min_month = subset['운용월'].min()
            max_month = subset['운용월'].max()
            self.doc.add_paragraph(f'- {dataset}: {min_month} ~ {max_month} 개월', style='List Bullet')
            
        self.doc.add_page_break()

        # --- 2. Methodology ---
        self._add_methodology_chapter()
        self.doc.add_page_break()

        # --- 3. Screening Summary ---
        self.doc.add_heading('3. 전체 항목 스크리닝 (Screening Summary)', level=1)
        self.doc.add_paragraph('전체 항목에 대한 위험도 분석 결과 (Top 10):')
        
        top_10 = screening_df.head(10)
        cols_to_show = ['Item', 'Slope', 'R2', 'Var_Ratio', 'Norm_Slope']
        available_cols = [c for c in cols_to_show if c in top_10.columns]
        self._add_dataframe_to_table(top_10[available_cols])
        
        self.doc.add_page_break()

        # --- 4. Detailed Analysis ---
        if is_full_report:
            for i, data in enumerate(items_data):
                self._add_item_section(data, section_num=4, subsection_num=i+1)
                self.doc.add_page_break()
        else:
            self._add_item_section(items_data, section_num=4)
        
        return self._save_to_stream()

    def _add_toc(self):
        """Inserts a Table of Contents field code."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        
        paragraph.add_run("\n(목차를 보려면 이 부분을 클릭하고 F9를 누르거나, 우클릭하여 '필드 업데이트'를 선택하세요.)").italic = True

    def _add_methodology_chapter(self):
        """Adds the Methodology chapter with detailed explanations."""
        self.doc.add_heading('2. 분석 기법 및 방법론 (Methodology)', level=1)
        self.doc.add_paragraph('본 장에서는 미사일 신뢰도 분석에 사용된 통계적 기법과 추세 예측 모델에 대해 상세히 설명합니다.')
        
        # 2.1 Distribution Analysis
        self.doc.add_heading('2.1 데이터 분포 분석 (Distribution Analysis)', level=2)
        
        self.doc.add_heading('KDE (Kernel Density Estimation)', level=3)
        self.doc.add_paragraph('KDE는 관측된 데이터로부터 확률 밀도 함수(PDF)를 추정하는 비모수적 방법입니다. 히스토그램보다 부드러운 곡선 형태로 데이터의 분포를 시각화하여, 데이터의 중심 경향성, 퍼짐 정도, 다봉성(Multi-modality) 등을 파악하는 데 유용합니다.')
        
        self.doc.add_heading('Box Plot (상자 그림)', level=3)
        self.doc.add_paragraph('Box Plot은 데이터의 다섯 수치 요약(최솟값, 제1사분위수, 중앙값, 제3사분위수, 최댓값)을 시각적으로 표현합니다. 이를 통해 데이터의 치우침 정도와 이상치(Outlier)를 직관적으로 식별할 수 있습니다.')
        
        # 2.2 Trend Prediction Models
        self.doc.add_heading('2.2 추세 예측 모델 (Trend Prediction Models)', level=2)
        self.doc.add_paragraph('저장 기간에 따른 특성치 변화를 예측하기 위해 다음과 같은 다양한 회귀 분석 모델을 활용합니다.')
        
        # Linear Regression
        self.doc.add_heading('1. Linear Regression (선형 회귀)', level=3)
        self.doc.add_paragraph('가장 기본적인 회귀 모델로, 독립 변수(운용월)와 종속 변수(특성치) 간의 관계를 직선으로 모델링합니다.')
        self.doc.add_paragraph('수식: y = w0 + w1 * x')
        self.doc.add_paragraph('여기서 y는 예측값, x는 운용월, w0는 절편, w1은 기울기를 의미합니다.')
        
        # Polynomial Regression
        self.doc.add_heading('2. Polynomial Regression (다항 회귀)', level=3)
        self.doc.add_paragraph('데이터가 비선형적인 경향을 보일 때 사용하며, 2차 이상의 다항식을 사용하여 곡선 형태의 추세를 적합합니다. 본 분석에서는 2차 다항식을 주로 사용합니다.')
        self.doc.add_paragraph('수식: y = w0 + w1 * x + w2 * x^2')
        
        # Bayesian Ridge
        self.doc.add_heading('3. Bayesian Ridge Regression (베이지안 릿지)', level=3)
        self.doc.add_paragraph('선형 회귀에 L2 규제(Regularization)를 적용하고, 베이지안 추론을 통해 회귀 계수의 확률 분포를 추정합니다. 데이터가 적거나 노이즈가 많을 때 과적합(Overfitting)을 방지하는 데 효과적입니다.')
        
        # Gaussian Process
        self.doc.add_heading('4. Gaussian Process Regressor (가우시안 프로세스)', level=3)
        self.doc.add_paragraph('커널(Kernel) 함수를 사용하여 데이터 간의 유사성을 기반으로 예측을 수행하는 비모수적 베이지안 방법입니다. 예측값뿐만 아니라 예측의 불확실성(신뢰 구간)을 함께 제공하여, 데이터가 없는 구간에서의 예측 신뢰도를 판단하는 데 유용합니다.')
        self.doc.add_paragraph('수식: f(x) ~ GP(m(x), k(x, x\'))')
        
        # SVR
        self.doc.add_heading('5. Support Vector Regression (SVR)', level=3)
        self.doc.add_paragraph('SVM(Support Vector Machine)의 회귀 버전으로, 마진(Margin) 내에 들어오는 오차는 허용하고 이를 벗어나는 오차를 최소화하는 초평면을 찾습니다. 커널 트릭을 사용하여 비선형 관계를 효과적으로 모델링할 수 있습니다.')
        
        # Neural Network
        self.doc.add_heading('6. Neural Network (인공신경망)', level=3)
        self.doc.add_paragraph('다층 퍼셉트론(MLP) 구조를 사용하여 입력과 출력 간의 복잡한 비선형 관계를 학습합니다. 은닉층(Hidden Layer)과 활성화 함수(Activation Function)를 통해 데이터의 내재된 패턴을 포착합니다.')

    def _add_item_section(self, data, section_num=3, subsection_num=None):
        item = data['item']
        stats_df = data['stats']
        figures = data['figures']
        model_metrics = data['metrics']
        
        header_text = f'{section_num}.{subsection_num} 상세 분석: Item {item}' if subsection_num else f'{section_num}. 상세 분석: Item {item}'
        self.doc.add_heading(header_text, level=1)
        
        # Statistics
        self.doc.add_heading('기초 통계량', level=2)
        # Reset index to show Dataset column
        stats_reset = stats_df.reset_index().rename(columns={'index': 'Dataset'})
        self._add_dataframe_to_table(stats_reset)
        
        # Distribution
        self.doc.add_heading('데이터 분포 (Distribution)', level=2)
        self.doc.add_paragraph("설명: KDE(Kernel Density Estimation) 그래프는 데이터의 확률 밀도 함수를 추정하여 분포의 형상을 시각화한 것입니다. Box Plot은 데이터의 중앙값, 사분위수, 이상치를 요약하여 보여줍니다.")
        
        if 'distribution_plot' in figures:
            self.doc.add_picture(figures['distribution_plot'], width=Inches(6))
            self.doc.add_paragraph(f'그림 {item}-1. 데이터 분포 비교 (KDE)', style='Caption')
            
        if 'box_plot' in figures:
            self.doc.add_picture(figures['box_plot'], width=Inches(6))
            self.doc.add_paragraph(f'그림 {item}-2. 데이터 분포 요약 (Box Plot)', style='Caption')
            
        # Interpretation for Distribution
        qim_mean = stats_df.loc['QIM', 'mean'] if 'QIM' in stats_df.index else 0
        asrp_mean = stats_df.loc['ASRP', 'mean'] if 'ASRP' in stats_df.index else 0
        diff = asrp_mean - qim_mean
        
        interp_text = f"해석: QIM(초기) 대비 ASRP(저장) 단계에서 평균값은 약 {diff:.4f} 만큼 {'증가' if diff > 0 else '감소'}했습니다."
        self.doc.add_paragraph(interp_text)
            
        # Trend Prediction
        self.doc.add_heading('추세 예측 (Trend Prediction)', level=2)
        self.doc.add_paragraph("설명: 다양한 회귀 모델을 사용하여 운용 기간(월)에 따른 측정값의 변화 추세를 예측한 결과입니다. 240개월(20년)까지의 장기 거동을 예측합니다.")
        
        if 'trend_plot' in figures:
            self.doc.add_picture(figures['trend_plot'], width=Inches(6))
            self.doc.add_paragraph(f'그림 {item}-3. 다중 모델 추세 예측 결과', style='Caption')
            
        # RMSE Table
        self.doc.add_heading('모델 성능 평가 (RMSE)', level=3)
        if model_metrics:
            metrics_df = pd.DataFrame(list(model_metrics.items()), columns=['Model', 'RMSE'])
            metrics_df = metrics_df.sort_values('RMSE')
            self._add_dataframe_to_table(metrics_df)
            
            best_model = metrics_df.iloc[0]
            self.doc.add_paragraph(f'\n해석: 현재 데이터에 대해 가장 적합한 모델은 "{best_model["Model"]}"이며, RMSE는 {best_model["RMSE"]:.4f}입니다. RMSE가 낮을수록 예측 정확도가 높음을 의미합니다.')

    def _add_dataframe_to_table(self, df):
        """Helper to convert DataFrame to Word Table."""
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            
        # Rows
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)
                    
    def _save_to_stream(self):
        """Saves the document to a BytesIO stream."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
